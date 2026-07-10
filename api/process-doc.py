from http.server import BaseHTTPRequestHandler
import json
import base64
import io
import traceback

try:
	from docx import Document
	from docx.enum.text import WD_ALIGN_PARAGRAPH
	DOCX_AVAILABLE = True
except ImportError:
	DOCX_AVAILABLE = False

try:
	from api.pii_scanner import scan_ir_for_pii, build_error_response, log_audit_event
except ImportError:
	try:
		from pii_scanner import scan_ir_for_pii, build_error_response, log_audit_event
	except ImportError:
		scan_ir_for_pii = None
		build_error_response = None
		log_audit_event = None

try:
	from api.docx_to_pdf import try_convert_docx_to_pdf
except ImportError:
	try:
		from docx_to_pdf import try_convert_docx_to_pdf
	except ImportError:
		try_convert_docx_to_pdf = None

try:
	from api.layout_raster import try_pdf_first_page_png, try_pdf_all_pages_png_list
except ImportError:
	try:
		from layout_raster import try_pdf_first_page_png, try_pdf_all_pages_png_list
	except ImportError:
		try_pdf_first_page_png = None
		try_pdf_all_pages_png_list = None

try:
	from api.strip_docx_fonts import strip_embedded_docx_fonts
except ImportError:
	try:
		from strip_docx_fonts import strip_embedded_docx_fonts
	except ImportError:
		strip_embedded_docx_fonts = None


def _align_to_str(alignment):
	if alignment == WD_ALIGN_PARAGRAPH.CENTER:
		return 'center'
	if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
		return 'right'
	if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
		return 'justify'
	return 'left'


def _extract_runs(paragraph):
	runs = []
	import re
	from docx.oxml.ns import qn
	
	# Extract ALL runs including those inside hyperlinks
	# paragraph.runs misses hyperlink content because w:hyperlink wraps w:r elements
	p_element = paragraph._p
	
	for child in p_element.iterchildren():
		tag = child.tag.rsplit('}', 1)[-1]
		
		if tag == 'r':
			# Normal run
			_append_run_from_element(child, runs, paragraph)
		elif tag == 'hyperlink':
			# Hyperlink - extract runs inside it, mark as underline (it's a link)
			for sub_run in child.iterchildren():
				sub_tag = sub_run.tag.rsplit('}', 1)[-1]
				if sub_tag == 'r':
					_append_run_from_element(sub_run, runs, paragraph, is_hyperlink=True)
	
	return runs


def _append_run_from_element(r_element, runs, paragraph, is_hyperlink=False):
	"""Extract run data from a w:r XML element and append to runs list."""
	from docx.oxml.ns import qn
	
	# Get text from all w:t elements in this run
	text_parts = []
	for t_elem in r_element.iterchildren():
		t_tag = t_elem.tag.rsplit('}', 1)[-1]
		if t_tag == 't':
			text_parts.append(t_elem.text or '')
	
	text = ''.join(text_parts)
	if not text:
		return
	
	# Extract formatting from rPr (run properties)
	rPr = r_element.find(qn('w:rPr'))
	
	is_bold = False
	is_italic = False
	is_underline = is_hyperlink  # Hyperlinks are treated as underlined
	font_size_pt = None
	font_name = None
	font_color = None      # hex string without leading '#', e.g. 'FFFFFF' or '2D73B5'
	highlight = None       # named highlight, e.g. 'yellow'
	run_shading_fill = None # hex string for run-level w:shd (banner-style colored heading)
	
	if rPr is not None:
		# Bold: <w:b/> or <w:b w:val="true"/>
		b_elem = rPr.find(qn('w:b'))
		if b_elem is not None:
			val = b_elem.get(qn('w:val'))
			is_bold = val is None or val.lower() in ('true', '1', 'on')
		
		# Italic: <w:i/>
		i_elem = rPr.find(qn('w:i'))
		if i_elem is not None:
			val = i_elem.get(qn('w:val'))
			is_italic = val is None or val.lower() in ('true', '1', 'on')
		
		# Underline: <w:u w:val="single"/>
		u_elem = rPr.find(qn('w:u'))
		if u_elem is not None:
			val = u_elem.get(qn('w:val'))
			if val and val.lower() != 'none':
				is_underline = True
		
		# Font size: <w:sz w:val="22"/> (half-points)
		sz_elem = rPr.find(qn('w:sz'))
		if sz_elem is not None:
			try:
				half_pts = int(sz_elem.get(qn('w:val')))
				font_size_pt = half_pts / 2.0
			except (ValueError, TypeError):
				pass
		
		# Font name: <w:rFonts w:ascii="Arial"/>
		rFonts = rPr.find(qn('w:rFonts'))
		if rFonts is not None:
			font_name = rFonts.get(qn('w:ascii'))
		
		# Font color: <w:color w:val="FFFFFF"/> — capture only meaningful colors
		# (skip 'auto', black, and missing values; downstream uses this to render
		# white-on-banner text and other intentional color formatting).
		color_elem = rPr.find(qn('w:color'))
		if color_elem is not None:
			val = (color_elem.get(qn('w:val')) or '').strip().upper()
			if val and val not in ('AUTO', '000000'):
				font_color = val
		
		# Highlight: <w:highlight w:val="yellow"/>
		hi_elem = rPr.find(qn('w:highlight'))
		if hi_elem is not None:
			val = (hi_elem.get(qn('w:val')) or '').strip().lower()
			if val and val != 'none':
				highlight = val
		
		# Run shading: <w:shd w:fill="2D73B5"/> — Word uses this for inline banners
		# (solid colored headings where the entire heading paragraph's runs share a fill).
		shd_elem = rPr.find(qn('w:shd'))
		if shd_elem is not None:
			fill = (shd_elem.get(qn('w:fill')) or '').strip().upper()
			if fill and fill not in ('AUTO', 'FFFFFF', ''):
				run_shading_fill = fill
	
	run = {
		'text': text,
		'bold': is_bold,
		'italic': is_italic,
		'underline': is_underline,
		'fontSizePt': font_size_pt,
		'fontFamily': font_name,
		'isHyperlink': is_hyperlink,
	}
	# Only attach color/highlight/shading if present, to keep IR compact for back-compat
	if font_color:
		run['fontColor'] = font_color
	if highlight:
		run['highlight'] = highlight
	if run_shading_fill:
		run['shadingFill'] = run_shading_fill
	runs.append(run)


def _detect_list_info(paragraph):
	"""
	Attempts to detect if a paragraph is part of a list and its level.
	python-docx does not expose numbering API directly; we inspect XML lightly.
	"""
	is_list = False
	level = None
	marker = None
	try:
		p = paragraph._p
		pPr = p.pPr
		if pPr is not None and pPr.numPr is not None:
			is_list = True
			# ilvl is the list indentation level
			if pPr.numPr.ilvl is not None:
				try:
					level = int(pPr.numPr.ilvl.val)
				except Exception:
					level = 0
			else:
				level = 0
			# We cannot know exact marker from Word reliably; leave None
			marker = None
	except Exception:
		pass
	return is_list, level, marker


def _extract_paragraph_ir(paragraph):
	from docx.oxml.ns import qn
	runs = _extract_runs(paragraph)
	full_text = ''.join(r.get('text') or '' for r in runs)

	# leading spaces count
	leading_spaces = 0
	for ch in full_text:
		if ch == ' ':
			leading_spaces += 1
		else:
			break

	is_list, level, marker = _detect_list_info(paragraph)

	# Detect paragraph border (bottom border = horizontal rule separator),
	# paragraph shading fill (background-color for banner divs), and page-break
	# markers.
	#
	# Word encodes page breaks two ways:
	#   1. pPr/w:pageBreakBefore  — the paragraph itself starts on a new page.
	#      Downstream: emit <hr> BEFORE this paragraph → pageBreakBefore=True.
	#   2. run-level w:br w:type="page"  — the break lives inside a run, and
	#      everything AFTER the break renders on the next page. In practice
	#      Word almost always places this in the LAST run of the previous
	#      paragraph, so the paragraph that follows in body order is what
	#      visually starts the new page. We record pageBreakAfter=True on
	#      THIS paragraph, and _build_ir_document promotes that flag to
	#      pageBreakBefore=True on the next block after all blocks are built.
	para_border_bottom = None
	para_shading_fill = None
	page_break_before = False
	page_break_after = False
	try:
		pPr_elem = paragraph._p.find(qn('w:pPr'))
		if pPr_elem is not None:
			pBdr = pPr_elem.find(qn('w:pBdr'))
			if pBdr is not None:
				bottom_el = pBdr.find(qn('w:bottom'))
				if bottom_el is not None:
					val = bottom_el.get(qn('w:val'), '')
					if val not in ('nil', 'none', ''):
						para_border_bottom = True
			# w:shd w:fill="DEEAF6" → light blue banner background
			shd_el = pPr_elem.find(qn('w:shd'))
			if shd_el is not None:
				fill = (shd_el.get(qn('w:fill')) or '').strip().upper()
				if fill and fill not in ('AUTO', 'FFFFFF', ''):
					para_shading_fill = fill
			# Page break before this paragraph: <w:pageBreakBefore/>
			pbb = pPr_elem.find(qn('w:pageBreakBefore'))
			if pbb is not None:
				val = pbb.get(qn('w:val'))
				# Present with no val, or truthy val → active page-break-before
				if val is None or val.lower() in ('true', '1', 'on'):
					page_break_before = True
		# Detect a run-level page break anywhere in this paragraph.
		for r_elem in paragraph._p.iter(qn('w:r')):
			for br_elem in r_elem.iter(qn('w:br')):
				br_type = br_elem.get(qn('w:type')) or ''
				if br_type.lower() == 'page':
					page_break_after = True
					break
			if page_break_after:
				break
	except Exception:
		pass

	para_ir = {
		'type': 'paragraph',
		'runs': runs,
		'align': _align_to_str(paragraph.paragraph_format.alignment),
		'leadingSpaces': leading_spaces if leading_spaces > 0 else None,
		'styleName': paragraph.style.name if getattr(paragraph, 'style', None) else None,
		'isListItem': is_list,
		'listLevel': level,
		'listMarker': marker,
		'spacingBeforePt': float(paragraph.paragraph_format.space_before.pt) if paragraph.paragraph_format.space_before else None,
		'spacingAfterPt': float(paragraph.paragraph_format.space_after.pt) if paragraph.paragraph_format.space_after else None,
		'lineHeightMultiple': None,
		'leftIndentPt': float(paragraph.paragraph_format.left_indent.pt) if getattr(paragraph.paragraph_format, 'left_indent', None) else None,
		'firstLineIndentPt': float(paragraph.paragraph_format.first_line_indent.pt) if getattr(paragraph.paragraph_format, 'first_line_indent', None) else None,
		'hangingIndentPt': None,
		'borderBottom': para_border_bottom,
	}
	if page_break_before:
		para_ir['pageBreakBefore'] = True
	if page_break_after:
		para_ir['pageBreakAfter'] = True
	if para_shading_fill:
		para_ir['shadingFill'] = para_shading_fill
	else:
		# Word often applies banner shading via run-level <w:shd> (rPr/shd) instead of
		# paragraph-level shading. If every non-whitespace run in this paragraph shares
		# the same shading fill, treat the whole paragraph as banner-shaded so downstream
		# prompt logic can emit a banner <div style="background-color: ...">.
		run_fills = {r.get('shadingFill') for r in runs if (r.get('text') or '').strip() and r.get('shadingFill')}
		non_shaded_text = [r for r in runs if (r.get('text') or '').strip() and not r.get('shadingFill')]
		if len(run_fills) == 1 and not non_shaded_text:
			para_ir['shadingFill'] = next(iter(run_fills))
	hanging = getattr(paragraph.paragraph_format, 'hanging_indent', None)
	if hanging is not None:
		try:
			para_ir['hangingIndentPt'] = float(getattr(hanging, 'pt', hanging))
		except Exception:
			para_ir['hangingIndentPt'] = None
	return para_ir


def _border_style(elem):
	"""Return a compact border descriptor string from a w:top/bottom/left/right element, or None."""
	if elem is None:
		return None
	val = elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') or ''
	sz  = elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz') or ''
	color = elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color') or ''
	if val in ('nil', 'none', ''):
		return 'none'
	parts = [val]
	if sz:
		try:
			parts.append(f'{int(sz)/8:.2g}pt')
		except Exception:
			parts.append(sz)
	if color and color.lower() not in ('auto', 'ffffff', ''):
		parts.append(f'#{color}')
	return ' '.join(parts)


def _table_border_summary(tblPr_elem):
	"""Summarise table-level border settings from w:tblPr XML element."""
	if tblPr_elem is None:
		return None
	ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
	tblBorders = tblPr_elem.find(f'{{{ns}}}tblBorders')
	if tblBorders is None:
		return None
	sides = {}
	for side in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
		el = tblBorders.find(f'{{{ns}}}{side}')
		s = _border_style(el)
		if s:
			sides[side] = s
	if not sides:
		return None
	# Classify: all outer = 'box', all none = 'none', has inner = 'grid', else 'mixed'
	outer = {sides.get(k) for k in ('top','bottom','left','right')}
	inner = {sides.get(k) for k in ('insideH','insideV')}
	outer_vis = all(v and v != 'none' for v in [sides.get(k) for k in ('top','bottom','left','right') if sides.get(k)])
	inner_vis = any(v and v != 'none' for v in [sides.get('insideH'), sides.get('insideV')] if v)
	has_any = any(v and v != 'none' for v in sides.values())
	if not has_any:
		kind = 'none'
	elif outer_vis and inner_vis:
		kind = 'grid'
	elif outer_vis and not inner_vis:
		kind = 'box'
	elif not outer_vis and inner_vis:
		kind = 'inner-only'
	else:
		kind = 'mixed'
	return {'kind': kind, 'sides': sides}


def _cell_border_summary(tcPr_elem):
	"""Summarise cell-level border overrides from w:tcPr element."""
	if tcPr_elem is None:
		return None
	ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
	tcBorders = tcPr_elem.find(f'{{{ns}}}tcBorders')
	if tcBorders is None:
		return None
	sides = {}
	for side in ('top', 'bottom', 'left', 'right'):
		el = tcBorders.find(f'{{{ns}}}{side}')
		s = _border_style(el)
		if s:
			sides[side] = s
	return sides if sides else None


def _tcw_pct(tcPr_elem, table_width_twips):
	"""Return approximate column width percentage from w:tcW, or None."""
	if tcPr_elem is None:
		return None
	ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
	tcW = tcPr_elem.find(f'{{{ns}}}tcW')
	if tcW is None:
		return None
	w_type = tcW.get(f'{{{ns}}}type') or ''
	w_val = tcW.get(f'{{{ns}}}w') or ''
	try:
		val = int(w_val)
	except Exception:
		return None
	if w_type == 'pct':
		return round(val / 50, 1)  # 50ths of a percent
	if w_type in ('dxa', '') and table_width_twips and table_width_twips > 0:
		return round(val / table_width_twips * 100, 1)
	return None


def _table_width_twips(tblPr_elem):
	if tblPr_elem is None:
		return None
	ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
	tblW = tblPr_elem.find(f'{{{ns}}}tblW')
	if tblW is None:
		return None
	try:
		return int(tblW.get(f'{{{ns}}}w') or '0') or None
	except Exception:
		return None


def _vmerge_info(tcPr_elem):
	"""Returns 'restart' (first merged cell), 'continue', or None."""
	if tcPr_elem is None:
		return None
	ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
	vm = tcPr_elem.find(f'{{{ns}}}vMerge')
	if vm is None:
		return None
	val = vm.get(f'{{{ns}}}val') or ''
	return 'restart' if val == 'restart' else 'continue'


def _gridspan(tcPr_elem):
	"""Return column span (int >= 1)."""
	if tcPr_elem is None:
		return 1
	ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
	gs = tcPr_elem.find(f'{{{ns}}}gridSpan')
	if gs is None:
		return 1
	try:
		return max(1, int(gs.get(f'{{{ns}}}val') or '1'))
	except Exception:
		return 1


def _cell_valign(tcPr_elem):
	if tcPr_elem is None:
		return None
	ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
	vAlign = tcPr_elem.find(f'{{{ns}}}vAlign')
	if vAlign is None:
		return None
	v = vAlign.get(f'{{{ns}}}val') or ''
	return v if v in ('top', 'center', 'bottom') else None


def _cell_shading_fill(tcPr_elem):
	"""Return cell background fill as hex (no '#'), or None for auto/white/missing."""
	if tcPr_elem is None:
		return None
	ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
	shd = tcPr_elem.find(f'{{{ns}}}shd')
	if shd is None:
		return None
	fill = (shd.get(f'{{{ns}}}fill') or '').strip().upper()
	if not fill or fill in ('AUTO', 'FFFFFF'):
		return None
	return fill


def _hex6(value):
	"""Normalize a color string (with optional '#', theme suffix like ' [660]', or 'rgb(...)') to 6-digit uppercase hex.
	Returns None for None/empty/non-hex inputs and ignores auto/white/black sentinels."""
	if not value:
		return None
	import re as _re
	s = str(value).strip()
	if s.startswith('#'):
		s = s[1:]
	m = _re.match(r'([0-9a-fA-F]{6})', s)
	if not m:
		return None
	hex_val = m.group(1).upper()
	if hex_val in ('FFFFFF',):
		return None
	return hex_val


def _shape_fill_for_txbx(txbx_elem):
	"""Walk up from a w:txbxContent element to find the enclosing shape (modern wsp or legacy VML)
	and return (shapeFillHex, shapeStrokeHex). Both are 6-digit uppercase hex or None."""
	A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
	WPS_NS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
	shape_fill = None
	shape_stroke = None
	parent = txbx_elem
	for _ in range(20):
		parent = parent.getparent()
		if parent is None:
			break
		tag = parent.tag.rsplit('}', 1)[-1] if isinstance(parent.tag, str) else ''
		# VML legacy shape (v:shape): fillcolor/strokecolor on the element itself
		if tag == 'shape':
			shape_fill = shape_fill or _hex6(parent.get('fillcolor'))
			shape_stroke = shape_stroke or _hex6(parent.get('strokecolor'))
		# Modern wps:wsp — look for child wps:spPr/a:solidFill/a:srgbClr
		if tag == 'wsp':
			spPr = parent.find(f'{{{WPS_NS}}}spPr')
			if spPr is not None:
				solidFill = spPr.find(f'{{{A_NS}}}solidFill')
				if solidFill is not None:
					srgb = solidFill.find(f'{{{A_NS}}}srgbClr')
					if srgb is not None and not shape_fill:
						shape_fill = _hex6(srgb.get('val'))
				ln = spPr.find(f'{{{A_NS}}}ln')
				if ln is not None:
					sf2 = ln.find(f'{{{A_NS}}}solidFill')
					if sf2 is not None:
						srgb = sf2.find(f'{{{A_NS}}}srgbClr')
						if srgb is not None and not shape_stroke:
							shape_stroke = _hex6(srgb.get('val'))
		if shape_fill and shape_stroke:
			break
	return shape_fill, shape_stroke


def _extract_table_ir(table):
	ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
	tblPr = table._tbl.find(f'{{{ns}}}tblPr')
	tbl_width = _table_width_twips(tblPr)
	tbl_borders = _table_border_summary(tblPr)
	tbl_style = None
	if tblPr is not None:
		ts = tblPr.find(f'{{{ns}}}tblStyle')
		if ts is not None:
			tbl_style = ts.get(f'{{{ns}}}val')

	rows_ir = []
	for row in table.rows:
		cells_ir = []
		# Build a deduplicated list of (cell_element, cell_object) pairs for this row
		seen_ids = set()
		row_cells = []
		for cell in row.cells:
			cid = id(cell._tc)
			if cid in seen_ids:
				continue
			seen_ids.add(cid)
			row_cells.append(cell)

		for cell in row_cells:
			tcPr = cell._tc.find(f'{{{ns}}}tcPr')
			content = []
			seen_para = set()
			for para in cell.paragraphs:
				key = id(para)
				if key in seen_para:
					continue
				seen_para.add(key)
				content.append(_extract_paragraph_ir(para))

			col_span = _gridspan(tcPr)
			vmerge = _vmerge_info(tcPr)
			cell_borders = _cell_border_summary(tcPr)
			valign = _cell_valign(tcPr)
			width_pct = _tcw_pct(tcPr, tbl_width)
			cell_shading = _cell_shading_fill(tcPr)

			cell_ir = {
				'content': content,
				'widthPct': width_pct,
				'align': None,
				'vAlign': valign,
				'colSpan': col_span if col_span > 1 else None,
				'vMerge': vmerge,
				'header': False,
				'borders': cell_borders,
			}
			if cell_shading:
				cell_ir['shadingFill'] = cell_shading
			cells_ir.append(cell_ir)
		rows_ir.append({'cells': cells_ir})

	return {
		'type': 'table',
		'rows': rows_ir,
		'widthPct': 100,
		'borderCollapse': True,
		'styleName': tbl_style,
		'tableBorders': tbl_borders,
	}


def _resolve_list_types(doc, blocks):
	"""Resolve whether list items are bullet or numbered by reading numbering.xml definitions."""
	import zipfile
	import re as _re
	
	bullet_chars = {'\uf0b7', '\u2022', '\u25cf', '\u25cb', '\u25a0', '\u25aa', '-', '\u2013', '\u2014'}
	
	try:
		numbering_xml = None
		temp = io.BytesIO()
		doc.save(temp)
		temp.seek(0)
		with zipfile.ZipFile(temp, 'r') as z:
			if 'word/numbering.xml' in z.namelist():
				with z.open('word/numbering.xml') as f:
					numbering_xml = f.read().decode('utf-8')
		
		if not numbering_xml:
			for b in blocks:
				if b.get('type') == 'paragraph' and b.get('isListItem'):
					b['listType'] = 'bullet'
			return
		
		num_to_abstract = {}
		for m in _re.finditer(r'<w:num\s+w:numId="(\d+)"[^>]*>.*?<w:abstractNumId\s+w:val="(\d+)"', numbering_xml, _re.DOTALL):
			num_to_abstract[m.group(1)] = m.group(2)
		
		abstract_fmt = {}
		# Also capture bullet font + char so downstream can pick the right glyph
		# (e.g. Wingdings U+F0FE = checkmark → {Symbol(ü)}; Wingdings U+F0A7 = square → {Symbol(§)}).
		abstract_bullet_font = {}   # abs_id -> font name (Wingdings, Symbol, Courier New, ...)
		abstract_bullet_char = {}   # abs_id -> raw bullet character (e.g. '\uf0fe')
		for m in _re.finditer(r'<w:abstractNum\s+w:abstractNumId="(\d+)"[^>]*>(.*?)</w:abstractNum>', numbering_xml, _re.DOTALL):
			abs_id = m.group(1)
			body = m.group(2)
			lvl_match = _re.search(r'<w:lvl\s+w:ilvl="0"[^>]*>(.*?)</w:lvl>', body, _re.DOTALL)
			if lvl_match:
				lvl_body = lvl_match.group(1)
				fmt_match = _re.search(r'<w:numFmt\s+w:val="([^"]+)"', lvl_body)
				if fmt_match:
					abstract_fmt[abs_id] = fmt_match.group(1)
				txt_match = _re.search(r'<w:lvlText\s+w:val="([^"]*)"', lvl_body)
				if txt_match and txt_match.group(1) in bullet_chars:
					abstract_fmt[abs_id] = 'bullet'
				# Bullet font lives in w:rPr > w:rFonts at this level
				fnt_match = (
					_re.search(r'<w:rFonts\s+[^/]*?w:ascii="([^"]+)"', lvl_body) or
					_re.search(r'<w:rFonts\s+[^/]*?w:hAnsi="([^"]+)"', lvl_body)
				)
				if fnt_match:
					abstract_bullet_font[abs_id] = fnt_match.group(1)
				if txt_match:
					abstract_bullet_char[abs_id] = txt_match.group(1)
		
		# Build a map from paragraph XML element to numId
		from docx.oxml.ns import qn
		para_numids = {}
		for b in blocks:
			if b.get('type') != 'paragraph' or not b.get('isListItem'):
				continue
			# We need the raw paragraph element to get numId
			# Since we don't store it, re-scan from doc paragraphs
		
		# Re-scan doc paragraphs to get numId mapping
		para_texts_to_numid = {}
		for para in doc.paragraphs:
			pPr = para._p.find(qn('w:pPr'))
			if pPr is not None:
				numPr = pPr.find(qn('w:numPr'))
				if numPr is not None:
					numId_elem = numPr.find(qn('w:numId'))
					if numId_elem is not None:
						num_id = numId_elem.get(qn('w:val'), '')
						text = ''.join(run.text for run in para.runs if run.text)
						para_texts_to_numid[text.strip()] = num_id
		
		for b in blocks:
			if b.get('type') != 'paragraph' or not b.get('isListItem'):
				continue
			text = ''.join(r.get('text', '') for r in b.get('runs', [])).strip()
			num_id = para_texts_to_numid.get(text, '')
			if num_id and num_id in num_to_abstract:
				abs_id = num_to_abstract[num_id]
				fmt = abstract_fmt.get(abs_id, '')
				if fmt == 'bullet':
					b['listType'] = 'bullet'
				elif fmt in ('decimal', 'lowerLetter', 'upperLetter', 'lowerRoman', 'upperRoman'):
					b['listType'] = 'numbered'
				else:
					b['listType'] = 'bullet'
				# Attach bullet font + raw char (used by generate-template to pick
				# the right glyph: Wingdings/Symbol → {Symbol(x)}; Courier/Arial → literal)
				if abs_id in abstract_bullet_font:
					b['bulletFont'] = abstract_bullet_font[abs_id]
				if abs_id in abstract_bullet_char:
					b['bulletChar'] = abstract_bullet_char[abs_id]
			else:
				b['listType'] = 'bullet'
	except Exception:
		for b in blocks:
			if b.get('type') == 'paragraph' and b.get('isListItem'):
				b['listType'] = 'bullet'


def _build_ir_document(doc):
	# CRITICAL: Accept all tracked changes first
	# Documents with track changes have content in <w:ins> tags that python-docx doesn't read
	from docx.oxml.ns import qn
	
	# Remove all deletions and unwrap insertions
	for element in list(doc.element.body.iter()):
		if element.tag == qn('w:del'):
			element.getparent().remove(element)
		elif element.tag == qn('w:ins'):
			parent = element.getparent()
			index = list(parent).index(element)
			for child in list(element):
				parent.insert(index, child)
				index += 1
			parent.remove(element)
	
	# Save and reload to ensure python-docx re-parses
	temp_bytes = io.BytesIO()
	doc.save(temp_bytes)
	temp_bytes.seek(0)
	doc = Document(temp_bytes)
	
	blocks = []
	# Extract headers first (for NMLID detection)
	# Headers are in doc.sections[].header (and first_page_header, even_page_header, etc.)
	header_texts = []
	try:
		for section in doc.sections:
			# Check all header types: default header, first page header, even page header
			headers_to_check = []
			if hasattr(section, 'header'):
				headers_to_check.append(section.header)
			if hasattr(section, 'first_page_header'):
				headers_to_check.append(section.first_page_header)
			if hasattr(section, 'even_page_header'):
				headers_to_check.append(section.even_page_header)
			
			for header in headers_to_check:
				try:
					for para in header.paragraphs:
						text = ''.join(run.text for run in para.runs if run.text)
						if text.strip():
							header_texts.append(text)
				except Exception:
					continue
	except Exception:
		pass  # Headers might not be accessible, continue without them
	
	# Extract text box content first (floating text boxes are not in body flow).
	# They appear as w:txbxContent inside drawing/shape elements. Word writes both
	# a modern (wsp/wps) AND a VML fallback (v:shape/v:textbox) for each textbox via
	# mc:AlternateContent — the two are visually identical but only the VML fallback
	# carries fillcolor/strokecolor as attributes; the modern one stores them in
	# child wps:spPr/a:solidFill nodes. We process BOTH variants, then dedup by
	# joined text content preferring whichever variant carries shape colors.
	#
	# Crucially we also capture the textbox's ANCHOR (the body-level w:p or w:tbl
	# that the floating shape is anchored to) so downstream rendering can place
	# the box at its actual source position instead of guessing from content.
	body_elem = doc.element.body
	body_children = list(body_elem.iterchildren())
	body_child_index = {id(el): idx for idx, el in enumerate(body_children)}

	def _body_anchor_for(txbx_elem):
		"""Walk up from a w:txbxContent until we hit an element whose parent IS the
		document body. Returns (anchor_element, body_index) or (None, None) if not
		found (e.g. txbx is inside a nested doc or header). The anchor is the
		body-level w:p or w:tbl in whose flow the floating shape lives."""
		n = txbx_elem
		while n is not None:
			parent = n.getparent()
			if parent is body_elem:
				idx = body_child_index.get(id(n))
				return n, idx
			n = parent
		return None, None

	raw_textboxes = []
	for txbx in body_elem.iter(qn('w:txbxContent')):
		rows = []
		seen_in_this_box = set()
		for p_elem in txbx.iter(qn('w:p')):
			text = ''.join(t.text or '' for t in p_elem.iter(qn('w:t')))
			text = text.strip()
			if not text:
				continue
			# Dedup ONLY within this textbox (not across textboxes — we need the
			# VML-fallback copy intact so we can read its fillcolor on the shape)
			if text in seen_in_this_box:
				continue
			seen_in_this_box.add(text)
			runs = []
			for r_elem in p_elem.iter(qn('w:r')):
				t_text = ''.join(t.text or '' for t in r_elem.iter(qn('w:t')))
				if not t_text:
					continue
				rPr = r_elem.find(qn('w:rPr'))
				is_bold = False
				if rPr is not None:
					b = rPr.find(qn('w:b'))
					if b is not None:
						val = b.get(qn('w:val'))
						is_bold = val is None or val.lower() in ('true', '1', 'on')
				runs.append({'text': t_text, 'bold': is_bold, 'italic': False, 'underline': False, 'fontSizePt': None, 'fontFamily': None, 'isHyperlink': False})
			if runs:
				rows.append({
					'type': 'paragraph',
					'runs': runs,
					'align': 'left',
					'leadingSpaces': None,
					'styleName': None,
					'isListItem': False,
					'listLevel': None,
					'listMarker': None,
					'spacingBeforePt': None,
					'spacingAfterPt': None,
					'lineHeightMultiple': None,
					'leftIndentPt': None,
					'firstLineIndentPt': None,
					'hangingIndentPt': None
				})
		if not rows:
			continue
		shape_fill, shape_stroke = _shape_fill_for_txbx(txbx)
		anchor_el, anchor_body_idx = _body_anchor_for(txbx)
		tb = {'type': 'textbox', 'rows': rows}
		if shape_fill:
			tb['shadingFill'] = shape_fill
		if shape_stroke:
			tb['borderColor'] = shape_stroke
		if anchor_el is not None:
			tb['_anchorElement'] = anchor_el  # resolved to block index after body iter
		if anchor_body_idx is not None:
			tb['_anchorBodyIndex'] = anchor_body_idx
		raw_textboxes.append(tb)
	# Dedup across textbox variants by joined text content. Prefer the variant that
	# carries shadingFill / borderColor — this is almost always the VML fallback.
	# We dedup ONLY across variants that share the same anchor body index, so
	# distinct textboxes that happen to carry identical text (which legitimately
	# happens in some letter families) are preserved as separate boxes.
	dedup = {}  # (text_key, anchor_idx_or_None) -> textbox
	for tb in raw_textboxes:
		text_key = '|'.join(
			''.join((r.get('text') or '') for r in row.get('runs', [])).strip()
			for row in tb.get('rows', [])
		)
		key = (text_key, tb.get('_anchorBodyIndex'))
		if key not in dedup:
			dedup[key] = tb
		else:
			existing = dedup[key]
			tb_has_colors = bool(tb.get('shadingFill') or tb.get('borderColor'))
			ex_has_colors = bool(existing.get('shadingFill') or existing.get('borderColor'))
			if tb_has_colors and not ex_has_colors:
				dedup[key] = tb
			elif tb_has_colors and ex_has_colors:
				for fld in ('shadingFill', 'borderColor'):
					if not existing.get(fld) and tb.get(fld):
						existing[fld] = tb[fld]
	# Preserve document order (body anchor index ascending; un-anchored last)
	text_box_blocks = sorted(
		dedup.values(),
		key=lambda x: (x.get('_anchorBodyIndex') if x.get('_anchorBodyIndex') is not None else 10**9)
	)

	# Iterate block items in document order: paragraphs and tables.
	# Track a mapping from the body-level element id → block index so we can
	# later resolve each textbox's anchor to a concrete block position.
	body_elem_to_block_index = {}
	for element in body_elem.iterchildren():
		tag = element.tag.rsplit('}', 1)[-1]
		if tag == 'p':
			for para in doc.paragraphs:
				if para._p is element:
					body_elem_to_block_index[id(element)] = len(blocks)
					blocks.append(_extract_paragraph_ir(para))
					break
		elif tag == 'tbl':
			for tbl in doc.tables:
				if tbl._tbl is element:
					body_elem_to_block_index[id(element)] = len(blocks)
					blocks.append(_extract_table_ir(tbl))
					break

	# Promote each `pageBreakAfter` flag onto the NEXT block as `pageBreakBefore`.
	# Word encodes a page break as `<w:br w:type="page"/>` inside the LAST run of
	# the paragraph that ends the current page — the paragraph that visually
	# starts the new page is the following body block. The downstream prompt
	# instructions treat `pageBreakBefore=True` as "emit <br>\n<hr>" immediately
	# before this block, which gives the on-screen HTML the same page boundary
	# the printed document has.
	for i in range(len(blocks) - 1):
		if isinstance(blocks[i], dict) and blocks[i].get('pageBreakAfter'):
			nxt = blocks[i + 1]
			if isinstance(nxt, dict):
				nxt['pageBreakBefore'] = True
			# Clear the raw pageBreakAfter marker so it doesn't confuse consumers
			# that only understand pageBreakBefore.
			blocks[i].pop('pageBreakAfter', None)

	# Resolve textbox anchors to block indices and surrounding-paragraph
	# snippets so downstream rendering can place each box at its true source
	# position (above/below whatever real paragraphs visually flank it). We
	# also normalize the IR by stripping the internal _anchorElement reference
	# now that we've used it.
	def _block_visible_text(block, max_len=120):
		if not isinstance(block, dict):
			return ''
		btype = block.get('type')
		text = ''
		if btype == 'paragraph':
			runs = block.get('runs') or []
			text = ''.join((r.get('text') or '') for r in runs).strip()
		elif btype == 'table':
			# Concat first row's cell texts as a coarse table preview
			rows = block.get('rows') or []
			if rows:
				cells = rows[0].get('cells') or []
				parts = []
				for c in cells:
					for p in (c.get('content') or []):
						for r in (p.get('runs') or []):
							t = (r.get('text') or '').strip()
							if t:
								parts.append(t)
				text = ' | '.join(parts)[:max_len]
		# Collapse whitespace
		text = ' '.join(text.split())
		return text[:max_len]

	for tb in text_box_blocks:
		anchor_el = tb.pop('_anchorElement', None)
		tb.pop('_anchorBodyIndex', None)
		if anchor_el is None:
			continue
		anchor_block_idx = body_elem_to_block_index.get(id(anchor_el))
		if anchor_block_idx is None:
			continue
		tb['anchorBlockIndex'] = anchor_block_idx
		# Walk backwards from the anchor to find the nearest non-empty preceding
		# block (skipping the anchor paragraph itself, which usually only holds
		# the floating shape and no real visible text).
		before_text = ''
		for i in range(anchor_block_idx - 1, -1, -1):
			t = _block_visible_text(blocks[i])
			if t:
				before_text = t
				break
		after_text = ''
		for i in range(anchor_block_idx, len(blocks)):
			# Skip the anchor block itself if it is empty / only carries the shape
			t = _block_visible_text(blocks[i])
			if t and i != anchor_block_idx:
				after_text = t
				break
			if t and i == anchor_block_idx:
				# Anchor itself has real content (rare); record it as "after"
				after_text = t
				break
		if before_text:
			tb['anchorBeforeText'] = before_text
		if after_text:
			tb['anchorAfterText'] = after_text
	# Resolve list types (bullet vs numbered) from numbering definitions
	_resolve_list_types(doc, blocks)
	
	# Extract document-level default font and size
	# Check docDefaults first, then fall back to "Body Text" and "Normal" paragraph styles
	default_font = None
	default_font_size_pt = None
	try:
		styles_elem = doc.element.find(qn('w:styles'))
		if styles_elem is not None:
			# 1. Try docDefaults
			doc_defaults = styles_elem.find(qn('w:docDefaults'))
			if doc_defaults is not None:
				rPrDefault = doc_defaults.find('.//' + qn('w:rPrDefault'))
				if rPrDefault is not None:
					rPr = rPrDefault.find(qn('w:rPr'))
					if rPr is not None:
						rFonts = rPr.find(qn('w:rFonts'))
						if rFonts is not None:
							default_font = (
								rFonts.get(qn('w:ascii')) or
								rFonts.get(qn('w:hAnsi')) or
								rFonts.get(qn('w:cs'))
							)
						sz = rPr.find(qn('w:sz'))
						if sz is not None:
							val = sz.get(qn('w:val'))
							if val:
								try:
									default_font_size_pt = int(val) / 2
								except Exception:
									pass
			# 2. If docDefaults didn't give us font, check "Body Text" then "Normal" styles
			if not default_font:
				for style_name_check in ('Body Text', 'Normal'):
					for style_elem in styles_elem.findall(qn('w:style')):
						name_elem = style_elem.find(qn('w:name'))
						if name_elem is not None and name_elem.get(qn('w:val')) == style_name_check:
							rPr = style_elem.find('.//' + qn('w:rPr'))
							if rPr is not None:
								rFonts = rPr.find(qn('w:rFonts'))
								sz = rPr.find(qn('w:sz'))
								if rFonts is not None:
									f = (rFonts.get(qn('w:ascii')) or
										 rFonts.get(qn('w:hAnsi')) or
										 rFonts.get(qn('w:cs')))
									if f and f not in ('Times New Roman', 'Calibri'):
										default_font = f
								if sz is not None and not default_font_size_pt:
									val = sz.get(qn('w:val'))
									if val:
										try:
											default_font_size_pt = int(val) / 2
										except Exception:
											pass
							if default_font:
								break
					if default_font:
						break
	except Exception:
		pass

	# Store header texts in meta for header detection
	return {
		'blocks': blocks,
		'source': 'docx',
		'confidence': 1.0,
		'images': [],
		'meta': {
			'headerTexts': header_texts,
			'textBoxes': text_box_blocks,
			'defaultFont': default_font,
			'defaultFontSizePt': default_font_size_pt,
		}
	}


class handler(BaseHTTPRequestHandler):
	def do_POST(self):
		try:
			content_length = int(self.headers.get('Content-Length', '0'))
			post_data = self.rfile.read(content_length)

			data = json.loads(post_data.decode('utf-8') or '{}')
			file_data = data.get('fileData')
			file_name = data.get('fileName', 'document.docx')

			if not file_data:
				return self._send(400, {'success': False, 'error': 'No file data provided'})
			if not DOCX_AVAILABLE:
				return self._send(500, {'success': False, 'error': 'python-docx library not available'})

			file_bytes = base64.b64decode(file_data)
			if strip_embedded_docx_fonts is not None:
				file_bytes, fonts_stripped, bytes_saved = strip_embedded_docx_fonts(file_bytes)
				if fonts_stripped:
					print(
						f"strip_docx_fonts: removed embedded fonts from {file_name} "
						f"({bytes_saved:,} bytes saved)"
					)
			doc = Document(io.BytesIO(file_bytes))

			ir = _build_ir_document(doc)

			# PII early-gate: scan the IR before returning it to the client
			pii_scan = None
			if scan_ir_for_pii is not None:
				pii_scan = scan_ir_for_pii(ir)
				if pii_scan.has_pii or pii_scan.severity == 'BLOCKED':
					error_msg = build_error_response(pii_scan)
					if log_audit_event:
						log_audit_event('DOC_BLOCKED', file_name, pii_scan, error_msg[:120] if error_msg else '')
					return self._send(403, {
						'success': False,
						'error': error_msg or 'Document blocked by PII policy scanner.',
						'pii_scan': pii_scan.to_dict()
					})

			if log_audit_event:
				log_audit_event('DOC_PROCESSED', file_name, pii_scan, 'IR extraction successful')

			payload = {'success': True, 'fileName': file_name, 'ir': ir}

			# Optional layout PDF (for browser screenshot / multimodal compare). Runs after PII gate.
			# Advanced mode rasterizes ALL pages (not just page 1) so Claude can see
			# the full visual structure of multi-page letters.
			if data.get('includeLayoutPdf') and try_convert_docx_to_pdf is not None:
				pdf_bytes, pdf_err = try_convert_docx_to_pdf(file_bytes, file_name)
				if pdf_bytes is not None:
					payload['layoutPdfBase64'] = base64.b64encode(pdf_bytes).decode('ascii')
					payload['layoutPdfMime'] = 'application/pdf'
					if try_pdf_all_pages_png_list is not None:
						png_pages, pages_err = try_pdf_all_pages_png_list(pdf_bytes)
						if png_pages:
							payload['layoutPngPages'] = [
								base64.b64encode(p).decode('ascii') for p in png_pages
							]
							payload['layoutPngPageCount'] = len(png_pages)
							# Keep single-page fallback for backwards compat
							payload['layoutPngBase64'] = payload['layoutPngPages'][0]
						elif pages_err:
							payload['layoutPngError'] = pages_err
					elif try_pdf_first_page_png is not None:
						# Fallback: single page only
						png_bytes, png_err = try_pdf_first_page_png(pdf_bytes)
						if png_bytes is not None:
							payload['layoutPngBase64'] = base64.b64encode(png_bytes).decode('ascii')
							payload['layoutPngPages'] = [payload['layoutPngBase64']]
							payload['layoutPngPageCount'] = 1
						elif png_err:
							payload['layoutPngError'] = png_err
				else:
					payload['layoutPdfError'] = pdf_err or 'unknown conversion error'

			return self._send(200, payload)
		except Exception as e:
			err = {
				'success': False,
				'error': f'{str(e)}',
				'trace': traceback.format_exc()
			}
			return self._send(500, err)

	def do_OPTIONS(self):
		self.send_response(200)
		self.send_header('Access-Control-Allow-Origin', '*')
		self.send_header('Access-Control-Allow-Headers', 'Content-Type')
		self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
		self.end_headers()

	def _send(self, status, payload):
		self.send_response(status)
		self.send_header('Content-type', 'application/json')
		self.send_header('Access-Control-Allow-Origin', '*')
		self.send_header('Access-Control-Allow-Headers', 'Content-Type')
		self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
		self.end_headers()
		self.wfile.write(json.dumps(payload).encode('utf-8'))

