from docx import Document
from docx.oxml.ns import qn
from io import BytesIO

def accept_tracked_changes(doc):
    for element in list(doc.element.body.iter()):
        if element.tag == qn('w:del'):
            element.getparent().remove(element)
        elif element.tag == qn('w:ins'):
            parent = element.getparent()
            index = list(parent).index(element)
            for child in list(element):
                parent.insert(index, child)
                index += 1
            parent.remove(element)
    temp_bytes = BytesIO()
    doc.save(temp_bytes)
    temp_bytes.seek(0)
    return Document(temp_bytes)

doc = Document(r'C:\Users\jhamrick\Downloads\CL008 - Loss Mit Consult 44 Day - Commerce - V1.0 (1).docx')
doc = accept_tracked_changes(doc)

print(f"Total paragraphs: {len(doc.paragraphs)}")
# Print from paragraph 90 onward
for i, p in enumerate(doc.paragraphs):
    if i < 90: continue
    t = p.text.strip()
    bold_runs = [r for r in p.runs if r.bold and r.text.strip()]
    is_bold = len(bold_runs) > 0
    is_underline = any(r.underline for r in p.runs if r.text.strip())
    # Also check run-level bold for mixed runs
    run_details = []
    for r in p.runs:
        if r.text.strip():
            run_details.append(f"b={r.bold},u={r.underline}:'{r.text[:30]}'")
    flag = ''
    if is_bold: flag += ' [BOLD]'
    if is_underline: flag += ' [UL]'
    if t:
        print(f"[{i}]{flag} {repr(t[:110])}")
        if len(run_details) > 1:
            print(f"     runs: {run_details}")
    else:
        print(f"[{i}] (blank)")
