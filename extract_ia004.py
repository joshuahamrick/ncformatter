from docx import Document
from docx.oxml.ns import qn
from io import BytesIO

def accept_tracked_changes(doc):
    for element in list(doc.element.body.iter()):
        if element.tag == qn('w:del'):
            element.getparent().remove(element)
        elif element.tag == qn('w:ins'):
            parent = element.getparent()
            index = list(parent).index(element)
            for child in list(element):
                parent.insert(index, child)
                index += 1
            parent.remove(element)
    temp_bytes = BytesIO()
    doc.save(temp_bytes)
    temp_bytes.seek(0)
    return Document(temp_bytes)

doc = Document(r'C:\Users\jhamrick\Downloads\IA004 - FHA Coverage Term - FUB - V1.0.docx')
doc = accept_tracked_changes(doc)

print("=== TABLE ===")
table = doc.tables[0]
print(f"Rows: {len(table.rows)}, Cols: {len(table.columns)}")
for r_idx, row in enumerate(table.rows):
    print(f"Row {r_idx}:")
    for c_idx, cell in enumerate(row.cells):
        text = cell.text.strip()
        print(f"  [{c_idx}]: {repr(text[:80])}")

print()
print("=== PARAGRAPHS (non-blank) ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        print(f"[{i}] {repr(t[:120])}")
