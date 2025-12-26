import sys
import io
from docx import Document

def extract_tables(docx_path):
    doc = Document(docx_path)
    tables = []
    for table_idx, table in enumerate(doc.tables):
        table_data = []
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell_idx, cell in enumerate(row.cells):
                cell_text = '\n'.join([para.text for para in cell.paragraphs])
                row_data.append(cell_text)
            table_data.append(row_data)
        tables.append(table_data)
    return tables

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python extract_table.py <docx_path>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    try:
        tables = extract_tables(docx_path)
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
        for table_idx, table in enumerate(tables):
            print(f"\n=== TABLE {table_idx + 1} ===")
            for row_idx, row in enumerate(table):
                print(f"\nRow {row_idx + 1}:")
                for cell_idx, cell in enumerate(row):
                    print(f"  Cell {cell_idx + 1}: {repr(cell[:100])}")
    except Exception as e:
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
        print(f"Error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        sys.exit(1)


