"""Debug tracked changes processing"""
import sys
sys.path.insert(0, 'api')
import io
from docx import Document
from docx.oxml.ns import qn

doc_path = r"C:\Users\jhamrick\Downloads\CA030 - CA Initial Contact - VCI - V1.0.docx"

with open(doc_path, 'rb') as f:
    file_bytes = f.read()

# Load document
doc = Document(io.BytesIO(file_bytes))

print("BEFORE accepting changes:")
print(f"  Paragraphs: {len(doc.paragraphs)}")
print(f"  First 5 para texts: {[p.text[:50] for p in doc.paragraphs[:5] if p.text]}")

# Accept tracked changes
print("\nAccepting tracked changes...")

ins_count = 0
del_count = 0

for element in list(doc.element.body.iter()):
    if element.tag == qn('w:del'):
        element.getparent().remove(element)
        del_count += 1
    elif element.tag == qn('w:ins'):
        parent = element.getparent()
        index = list(parent).index(element)
        for child in list(element):  # Make a list copy
            parent.insert(index, child)
            index += 1
        parent.remove(element)
        ins_count += 1

print(f"  Removed {del_count} deletions")
print(f"  Unwrapped {ins_count} insertions")

# Save and reload
print("\nSaving and reloading...")
temp_bytes = io.BytesIO()
doc.save(temp_bytes)
print(f"  Saved bytes: {len(temp_bytes.getvalue())}")

temp_bytes.seek(0)
doc = Document(temp_bytes)

print("\nAFTER accepting changes:")
print(f"  Paragraphs: {len(doc.paragraphs)}")

# Show non-empty paragraphs
non_empty = [p.text for p in doc.paragraphs if p.text.strip()]
print(f"  Non-empty paragraphs: {len(non_empty)}")

for idx, text in enumerate(non_empty[:10]):
    print(f"    [{idx}] {text[:80]}")
