"""Deep dive into ALL paragraph runs and colors"""
import sys
sys.path.insert(0, 'api')

from docx import Document

doc_path = r"C:\Users\jhamrick\Downloads\CA030 - CA Initial Contact - VCI - V1.0.docx"
doc = Document(doc_path)

print("ANALYZING ALL PARAGRAPHS AND RUNS (including empty):\n")

for p_idx in range(min(50, len(doc.paragraphs))):  # First 50 paragraphs
    para = doc.paragraphs[p_idx]
    
    if len(para.runs) == 0:
        continue
    
    print(f"\n--- Paragraph {p_idx} ---")
    print(f"Para text: '{para.text[:80]}'")
    print(f"Runs: {len(para.runs)}")
    
    for r_idx, run in enumerate(para.runs):
        # Get all color properties
        color_str = "NO COLOR"
        try:
            if run.font.color:
                if run.font.color.rgb:
                    rgb = run.font.color.rgb
                    color_str = f"RGB({rgb[0]},{rgb[1]},{rgb[2]})"
                elif run.font.color.theme_color:
                    color_str = f"THEME_COLOR({run.font.color.theme_color})"
                else:
                    color_str = "COLOR_OBJ_NO_RGB"
        except Exception as e:
            color_str = f"ERROR: {e}"
        
        print(f"  Run {r_idx}: {color_str:25s} | Text: '{run.text[:60]}'")

print("\n\n--- LOOKING AT PARAGRAPHS 140-155 (near Sincerely) ---")
for p_idx in range(140, min(155, len(doc.paragraphs))):
    para = doc.paragraphs[p_idx]
    if para.text.strip() or len(para.runs) > 0:
        print(f"\nP{p_idx}: '{para.text}' ({len(para.runs)} runs)")
        for r_idx, run in enumerate(para.runs):
            if run.text:
                print(f"  R{r_idx}: '{run.text}'")
