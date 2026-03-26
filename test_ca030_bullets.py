#!/usr/bin/env python3
"""
Test CA030 content to see if bullet points have bullet characters
"""
import sys
import os
import io
from docx import Document
from docx.oxml.ns import qn

# Path to the CA030 document
doc_path = r"C:\Users\jhamrick\Downloads\CA030 - CA Initial Contact - VCI - V1.0.docx"

if not os.path.exists(doc_path):
    print(f"[ERROR] Document not found: {doc_path}")
    sys.exit(1)

print("Opening document...")
doc = Document(doc_path)

# Accept all tracked changes
print("Accepting tracked changes...")
for element in list(doc.element.body.iter()):
    if element.tag == qn('w:del'):
        element.getparent().remove(element)
    elif element.tag == qn('w:ins'):
        parent = element.getparent()
        index = list(parent).index(element)
        for child in list(element):
            parent.insert(index, child)
            index += 1
        parent.remove(element)

temp_bytes = io.BytesIO()
doc.save(temp_bytes)
temp_bytes.seek(0)
doc = Document(temp_bytes)

print("\n=== Looking for bullet points ===\n")

found_options = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # Find the "options" paragraph
    if 'options' in text.lower() and 'including' in text.lower():
        found_options = True
        print(f"Paragraph {i + 1}: '{text}'")
        print(f"\nNext 4 paragraphs:")
        continue
    
    if found_options and i < len(doc.paragraphs) - 1:
        # Check next 4 paragraphs
        for j in range(4):
            if i + j + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + j + 1]
                next_text = next_para.text.strip()
                if next_text:
                    # Check if it's a bullet point
                    is_list, level, marker = False, None, None
                    try:
                        p = next_para._p
                        pPr = p.pPr
                        if pPr is not None and pPr.numPr is not None:
                            is_list = True
                            if pPr.numPr.ilvl is not None:
                                level = int(pPr.numPr.ilvl.val)
                    except:
                        pass
                    
                    print(f"  Paragraph {i + j + 2}:")
                    print(f"    Text: '{next_text}'")
                    print(f"    Is List: {is_list}")
                    print(f"    Level: {level}")
                    print()
        break

print("\n=== Done ===")
