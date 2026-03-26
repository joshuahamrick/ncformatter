"""Test script to see what's being extracted from CA030 document"""
import sys
sys.path.insert(0, 'api')

from docx import Document
import json

# Read the CA030 document
doc_path = r"C:\Users\jhamrick\Downloads\CA030 - CA Initial Contact - VCI - V1.0.docx"
doc = Document(doc_path)

print("=" * 80)
print(f"TOTAL PARAGRAPHS: {len(doc.paragraphs)}")
print(f"TOTAL TABLES: {len(doc.tables)}")
print("=" * 80)

# Count non-empty paragraphs
non_empty = [p for p in doc.paragraphs if p.text.strip()]
print(f"NON-EMPTY PARAGRAPHS: {len(non_empty)}")

print("\n" + "=" * 80)
print("FIRST 30 NON-EMPTY PARAGRAPHS:")
print("=" * 80)

for i, para in enumerate(non_empty[:30]):
    text = para.text.strip()
    print(f"\n[{i}] {text[:150]}{'...' if len(text) > 150 else ''}")

print("\n" + "=" * 80)
print("LAST 15 NON-EMPTY PARAGRAPHS:")
print("=" * 80)

for i, para in enumerate(non_empty[-15:]):
    text = para.text.strip()
    print(f"\n[{len(non_empty)-15+i}] {text[:150]}{'...' if len(text) > 150 else ''}")
