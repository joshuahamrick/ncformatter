#!/usr/bin/env python3
"""
Test to see what labels are extracted from CA030 document
"""
import sys
import os
import io
from docx import Document
from docx.oxml.ns import qn

# Path to the CA030 document
doc_path = r"C:\Users\jhamrick\Downloads\CA030 - CA Initial Contact - VCI - V1.0.docx"

if not os.path.exists(doc_path):
    print(f"[ERROR] Document not found: {doc_path}")
    sys.exit(1)

print("Opening document...")
doc = Document(doc_path)

print("\n=== Accepting all tracked changes ===")
# Remove all deletions and unwrap insertions
for element in list(doc.element.body.iter()):
    if element.tag == qn('w:del'):
        element.getparent().remove(element)
    elif element.tag == qn('w:ins'):
        parent = element.getparent()
        index = list(parent).index(element)
        for child in list(element):
            parent.insert(index, child)
            index += 1
        parent.remove(element)

# Save and reload to ensure python-docx re-parses
temp_bytes = io.BytesIO()
doc.save(temp_bytes)
temp_bytes.seek(0)
doc = Document(temp_bytes)

print("\n=== Looking for Loan Number / RE / Property Address in paragraphs ===\n")

for i, para in enumerate(doc.paragraphs):
    text = para.text
    text_stripped = text.strip()
    
    if not text_stripped:
        continue
    
    # Look for relevant keywords
    text_lower = text_stripped.lower()
    if any(kw in text_lower for kw in ['loan number', 'property address', 're:', '{[m594]', '{[m567]']):
        leading_spaces = len(text) - len(text.lstrip(' '))
        
        print(f"Paragraph {i + 1}:")
        print(f"  Text: '{text_stripped}'")
        print(f"  Leading spaces: {leading_spaces}")
        print(f"  Has tabs: {chr(9) in text}")  # Check for tab character
        print()

print("\n=== Done ===")

