"""End-to-end test of CA030 document processing"""
import sys
sys.path.insert(0, 'api')

from docx import Document
import json

# Process the document with our color filtering
doc_path = r"C:\Users\jhamrick\Downloads\CA030 - CA Initial Contact - VCI - V1.0.docx"
doc = Document(doc_path)

print("=" * 80)
print("TESTING COLOR-BASED FILTERING")
print("=" * 80)

def extract_paragraph_with_color_filter(paragraph):
    """Extract paragraph, filtering out colored markup"""
    runs = []
    full_text = ''
    
    for run in paragraph.runs:
        skip_run = False
        
        try:
            # Check font color
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                r, g, b = rgb[0], rgb[1], rgb[2]
                if not (r < 50 and g < 50 and b < 50):
                    skip_run = True
                    print(f"  SKIPPING colored run: RGB({r},{g},{b}) - '{run.text[:50]}'")
            
            # Check highlighting
            if hasattr(run.font, 'highlight_color') and run.font.highlight_color:
                skip_run = True
                print(f"  SKIPPING highlighted run: '{run.text[:50]}'")
        except Exception as e:
            pass
        
        if not skip_run and run.text.strip():
            runs.append({
                'text': run.text,
                'bold': run.bold,
                'underline': run.underline
            })
            full_text += run.text
    
    return {
        'text': full_text,
        'runs': runs,
        'bold': any(r.get('bold') for r in runs),
        'underline': any(r.get('underline') for r in runs)
    }

# Extract all paragraphs
extracted = []
for i, para in enumerate(doc.paragraphs):
    para_data = extract_paragraph_with_color_filter(para)
    if para_data['text'].strip():
        extracted.append(para_data)
        if i < 30 or i > len(doc.paragraphs) - 10:  # First 30 and last 10
            print(f"\n[Para {i}] {para_data['text'][:100]}")

print(f"\n\n{'='*80}")
print(f"TOTAL EXTRACTED PARAGRAPHS: {len(extracted)}")
print(f"{'='*80}")

# Search for key content
print("\nSEARCHING FOR KEY CONTENT:")
key_phrases = ["Loan Number", "RE:", "Sincerely", "CompanyLongName", 
               "CompanyReturnAddr", "SPOCContactPhone", "FEDERAL LAW"]

for phrase in key_phrases:
    found = False
    for para in extracted:
        if phrase.lower() in para['text'].lower():
            print(f"\n✓ '{phrase}' FOUND: {para['text'][:80]}")
            found = True
            break
    if not found:
        print(f"\n✗ '{phrase}' NOT FOUND")
