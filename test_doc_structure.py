"""End-to-end test of CA030 document processing"""
import sys
sys.path.insert(0, 'api')

from docx import Document
import json

# Process the document with our color filtering
doc_path = r"C:\Users\jhamrick\Downloads\CA030 - CA Initial Contact - VCI - V1.0.docx"
doc = Document(doc_path)

print("=" * 80)
print("DOCUMENT STRUCTURE ANALYSIS")
print("=" * 80)
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")

# Check if content is in tables
if doc.tables:
    print("\nTABLES FOUND - Extracting table content:")
    for t_idx, table in enumerate(doc.tables):
        print(f"\n  Table {t_idx}: {len(table.rows)} rows x {len(table.columns)} cols")
        for r_idx, row in enumerate(table.rows[:10]):  # First 10 rows
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    print(f"    [{r_idx},{c_idx}] {text[:60]}")

# Check headers/footers
print("\nCHECKING HEADERS:")
for s_idx, section in enumerate(doc.sections):
    if hasattr(section, 'header'):
        for p_idx, para in enumerate(section.header.paragraphs):
            if para.text.strip():
                print(f"  Header section {s_idx}, para {p_idx}: {para.text[:60]}")

print("\nCHECKING ALL PARAGRAPH RUNS:")
non_empty_count = 0
for p_idx, para in enumerate(doc.paragraphs):
    if para.text.strip():
        non_empty_count += 1
        
    # Check all runs, including empty/colored
    for r_idx, run in enumerate(para.runs):
        if not run.text.strip():
            continue
            
        # Get color info
        color_info = "NO COLOR"
        try:
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                color_info = f"RGB({rgb[0]},{rgb[1]},{rgb[2]})"
        except:
            color_info = "COLOR ERROR"
        
        # Show first 20 and last 10 paragraphs with content
        if p_idx < 20 or p_idx > len(doc.paragraphs) - 10:
            print(f"  [P{p_idx:3d} R{r_idx}] {color_info:20s} | {run.text[:60]}")

print(f"\nNON-EMPTY PARAGRAPHS: {non_empty_count}")
