#!/usr/bin/env python3
"""
Test LM001 through the NcFormatter pipeline locally.
Outputs the generated HTML so we can compare to the gold standard.
"""
import json
import sys
import os
import importlib.util

# Module names have dashes so use importlib
def load_module(name, path):
    spec = importlib.util.spec_from_file_location(name, path)
    mod = importlib.util.module_from_spec(spec)
    sys.modules[name] = mod
    spec.loader.exec_module(mod)
    return mod

BASE = os.path.dirname(__file__)
API = os.path.join(BASE, 'api')

# Pre-load api sub-modules that process-doc and generate-template import
for sub in ['pii_scanner', 'docx_to_pdf', 'layout_raster', 'anthropic_retry']:
    path = os.path.join(API, f'{sub}.py')
    if os.path.exists(path):
        try:
            load_module(f'api.{sub}', path)
            load_module(sub, path)
        except Exception as e:
            print(f"  Warning: could not pre-load {sub}: {e}")

process_doc = load_module('process_doc', os.path.join(API, 'process-doc.py'))
gen_tmpl    = load_module('gen_tmpl',    os.path.join(API, 'generate-template.py'))

from docx import Document
import io

DOCX_PATH = r"K:\ncConnect\Letters\Triad\Letters\LM001\20231121\LM001 Loss Mit Acknowledgement Letter - Triad -V1.0.docx"

def main():
    print("=" * 70)
    print("LM001 Formatter Test")
    print("=" * 70)

    # Step 1: Load the docx
    print(f"\n[1] Loading DOCX...")
    with open(DOCX_PATH, 'rb') as f:
        file_bytes = f.read()

    doc = Document(io.BytesIO(file_bytes))
    print(f"    Document loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

    # Step 2: Build IR
    print("\n[2] Building IR...")
    ir = process_doc._build_ir_document(doc)
    blocks = ir.get('blocks', [])
    print(f"    IR built: {len(blocks)} blocks")

    with open(os.path.join(BASE, 'lm001_ir.json'), 'w', encoding='utf-8') as f:
        json.dump(ir, f, indent=2, default=str)
    print("    Saved: lm001_ir.json")

    # Step 3: Format IR for prompt
    print("\n[3] Formatting IR for prompt...")
    ir_content = gen_tmpl.format_ir_for_prompt(ir)

    with open(os.path.join(BASE, 'lm001_ir_formatted.txt'), 'w', encoding='utf-8') as f:
        f.write(ir_content)
    print(f"    Saved: lm001_ir_formatted.txt ({len(ir_content)} chars)")

    print("\n--- FORMATTED IR (first 4000 chars) ---")
    print(ir_content[:4000])
    if len(ir_content) > 4000:
        print(f"\n... [{len(ir_content)-4000} more chars] ...")
        print("\n--- LAST 2000 chars ---")
        print(ir_content[-2000:])

    # Step 4: Call Claude API
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        env_path = os.path.join(BASE, '.env')
        if os.path.exists(env_path):
            with open(env_path) as f:
                for line in f:
                    if line.startswith('ANTHROPIC_API_KEY='):
                        api_key = line.split('=', 1)[1].strip().strip('"\'')
                        os.environ['ANTHROPIC_API_KEY'] = api_key
                        break

    if not api_key:
        print("\n[4] No ANTHROPIC_API_KEY - skipping generation step")
        return

    print(f"\n[4] Calling Claude API...")
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)

    few_shot_examples = gen_tmpl.load_few_shot_examples()
    print(f"    Loaded {len(few_shot_examples)} examples: {[e['name'] for e in few_shot_examples]}")

    system_prompt, user_message, few_shot_text = gen_tmpl.build_prompt(ir, few_shot_examples)
    full_system = system_prompt + "\n\n" + few_shot_text

    print(f"    System: {len(full_system)} chars | User: {len(user_message)} chars")

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=10000,
        system=full_system,
        messages=[{"role": "user", "content": user_message}],
        temperature=0,
    )

    html = response.content[0].text.strip()
    if html.startswith('```html'):
        html = html.replace('```html', '').replace('```', '').strip()
    elif html.startswith('```'):
        html = html.replace('```', '').strip()

    html = gen_tmpl.normalize_html(html)

    out_path = os.path.join(BASE, 'lm001_generated.html')
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(html)
    print(f"\n[5] Generated HTML saved: lm001_generated.html ({len(html)} chars)")

    print("\n--- GENERATED HTML ---")
    print(html)

if __name__ == '__main__':
    main()
