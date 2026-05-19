#!/usr/bin/env python3
"""Quick test - build IR, format it, and check bullet char detection."""
import sys, os, importlib.util, json, io

def load_module(name, path):
    spec = importlib.util.spec_from_file_location(name, path)
    mod = importlib.util.module_from_spec(spec)
    sys.modules[name] = mod
    spec.loader.exec_module(mod)
    return mod

BASE = r'C:\Users\jhamrick\Desktop\NcFormatter'
API = os.path.join(BASE, 'api')

print('Loading modules...', flush=True)
for sub in ['pii_scanner', 'docx_to_pdf', 'layout_raster', 'anthropic_retry']:
    path = os.path.join(API, f'{sub}.py')
    if os.path.exists(path):
        try:
            load_module(f'api.{sub}', path)
            load_module(sub, path)
        except Exception as e:
            print(f'  Warning: {sub}: {e}', flush=True)

process_doc = load_module('process_doc', os.path.join(API, 'process-doc.py'))
gen_tmpl    = load_module('gen_tmpl',    os.path.join(API, 'generate-template.py'))
print('Modules loaded', flush=True)

from docx import Document
DOCX_PATH = r'K:\ncConnect\Letters\Triad\Letters\LM001\20231121\LM001 Loss Mit Acknowledgement Letter - Triad -V1.0.docx'
print('Loading docx...', flush=True)
with open(DOCX_PATH, 'rb') as f:
    file_bytes = f.read()
doc = Document(io.BytesIO(file_bytes))
print(f'  Doc: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables', flush=True)

print('\nBuilding IR...', flush=True)
ir = process_doc._build_ir_document(doc)
blocks = ir['blocks']
print(f'  IR: {len(blocks)} blocks', flush=True)

print('\nList items found (with bullet chars):', flush=True)
for i, b in enumerate(blocks):
    if b.get('isListItem'):
        runs = b.get('runs', [])
        text = ''.join(r.get('text','') for r in runs)[:55]
        bchar = b.get('listBulletChar', 'NOT SET')
        ltype = b.get('listType', 'NOT SET')
        print(f'  Block {i}: level={b.get("listLevel")} type={ltype} char={repr(bchar)} "{text}"', flush=True)

print('\nFormatting IR for prompt...', flush=True)
ir_text = gen_tmpl.format_ir_for_prompt(ir)
with open(os.path.join(BASE, 'lm001_ir_formatted.txt'), 'w', encoding='utf-8') as f:
    f.write(ir_text)
print(f'  Saved lm001_ir_formatted.txt ({len(ir_text)} chars)', flush=True)

print('\n--- FORMATTED IR (full) ---', flush=True)
print(ir_text, flush=True)
print('\nDone!', flush=True)
