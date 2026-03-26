"""Check if content is in shapes, text boxes, or other Word elements"""
import sys
sys.path.insert(0, 'api')

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

doc_path = r"C:\Users\jhamrick\Downloads\CA030 - CA Initial Contact - VCI - V1.0.docx"
doc = Document(doc_path)

print("CHECKING FOR CONTENT IN OTHER LOCATIONS:\n")

# Check if content is in document body's XML
print("1. Checking raw XML for text content...")
body_xml = doc._element.body
full_xml_text = body_xml.xml.decode('utf-8') if isinstance(body_xml.xml, bytes) else str(body_xml)

# Search for key phrases in raw XML
key_phrases = ["Loan Number", "M594", "M567", "CompanyLongName", "CompanyReturnAddr", 
               "FEDERAL LAW", "DEBT COLLECTOR", "SPOCContactPhone"]

for phrase in key_phrases:
    if phrase in full_xml_text:
        print(f"   FOUND '{phrase}' in XML!")
        # Find context
        idx = full_xml_text.find(phrase)
        context = full_xml_text[max(0, idx-100):min(len(full_xml_text), idx+200)]
        print(f"   Context: ...{context}...")
    else:
        print(f"   NOT found: {phrase}")

# Check document properties
print("\n2. Document Properties:")
print(f"   Core properties: {doc.core_properties}")
try:
    print(f"   Last modified by: {doc.core_properties.last_modified_by}")
    print(f"   Revision: {doc.core_properties.revision}")
except:
    pass

# Check sections
print(f"\n3. Document Sections: {len(doc.sections)}")
for idx, section in enumerate(doc.sections):
    print(f"   Section {idx}: {section.start_type}")
    
    # Check headers and footers
    if hasattr(section, 'header'):
        header_text = []
        for para in section.header.paragraphs:
            if para.text.strip():
                header_text.append(para.text.strip())
        if header_text:
            print(f"     Header: {header_text}")
    
    if hasattr(section, 'footer'):
        footer_text = []
        for para in section.footer.paragraphs:
            if para.text.strip():
                footer_text.append(para.text.strip())
        if footer_text:
            print(f"     Footer: {footer_text}")

# Check for text boxes or shapes
print("\n4. Checking for shapes/text boxes...")
# Text boxes are stored in the XML as w:txbxContent
if 'txbxContent' in full_xml_text:
    print("   TEXT BOXES FOUND in document!")
else:
    print("   No text boxes found")

print("\n5. First 500 chars of XML (looking for structure):")
print(full_xml_text[:500])
