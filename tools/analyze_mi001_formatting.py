#!/usr/bin/env python3
"""
Analyze MI001 document for bold formatting
"""
import sys
import os
import io
from docx import Document
from docx.oxml.ns import qn

doc_path = r"C:\Users\jhamrick\Downloads\MI001 - PMI Brrwr Can Req Inquiry -Keesler - V3.0.docx"

if not os.path.exists(doc_path):
    print(f"[ERROR] Document not found: {doc_path}")
    sys.exit(1)

print("Opening document...")
doc = Document(doc_path)

# Accept tracked changes
print("Accepting tracked changes...")
for element in list(doc.element.body.iter()):
    if element.tag == qn('w:del'):
        element.getparent().remove(element)
    elif element.tag == qn('w:ins'):
        parent = element.getparent()
        index = list(parent).index(element)
        for child in list(element):
            parent.insert(index, child)
            index += 1
        parent.remove(element)

temp_bytes = io.BytesIO()
doc.save(temp_bytes)
temp_bytes.seek(0)
doc = Document(temp_bytes)

print("\n=== Analyzing paragraphs for bold text ===\n")

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    
    # Check if paragraph has bold runs
    has_bold = False
    bold_text = []
    for run in para.runs:
        if run.bold and run.text.strip():
            has_bold = True
            bold_text.append(run.text)
    
    if has_bold:
        print(f"Paragraph {i + 1}:")
        print(f"  Full text: {text[:100]}...")
        print(f"  Bold parts: {', '.join([repr(t) for t in bold_text[:3]])}")
        print()

print("\n=== Done ===")
