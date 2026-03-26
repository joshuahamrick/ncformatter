import docx
from docx.oxml.ns import qn

doc = docx.Document(r'c:\Users\jhamrick\Downloads\- Addendum B - CS036 - Verification of Mortgage - Frost - V5.0.docx')

# Get ALL paragraphs with full run detail
print('=== ALL PARAGRAPHS WITH RUN DETAIL ===')
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if t.strip():
        print(f'\nP{i}: [{t}]')
        print(f'  Alignment: {p.alignment}')
        for ri, r in enumerate(p.runs):
            flags = []
            if r.bold: flags.append('BOLD')
            if r.underline: flags.append('UNDERLINE')
            if r.italic: flags.append('ITALIC')
            if r.font.size: flags.append(f'SIZE={r.font.size}')
            if r.font.name: flags.append(f'FONT={r.font.name}')
            flag_str = f' [{", ".join(flags)}]' if flags else ''
            print(f'  Run{ri}: [{r.text}]{flag_str}')

# Get table cell details with run formatting
print('\n\n=== TABLE CELL DETAIL WITH FORMATTING ===')
for ti, table in enumerate(doc.tables):
    print(f'\n--- Table {ti} ---')
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                print(f'\n  R{ri}C{ci}: [{text}]')
                for pi, p in enumerate(cell.paragraphs):
                    pt = p.text.strip()
                    if pt:
                        for rri, r in enumerate(p.runs):
                            flags = []
                            if r.bold: flags.append('BOLD')
                            if r.underline: flags.append('UNDERLINE')
                            flag_str = f' [{", ".join(flags)}]' if flags else ''
                            print(f'    P{pi}R{rri}: [{r.text}]{flag_str}')

# Check for any images
print('\n\n=== IMAGE CHECK ===')
for rel in doc.part.rels.values():
    if 'image' in rel.reltype:
        print(f'  Image: {rel.target_ref}')

# Paragraphs around the delinquent section
print('\n\n=== PARAGRAPHS 48-72 (around delinquent section) ===')
for i in range(48, min(73, len(doc.paragraphs))):
    p = doc.paragraphs[i]
    print(f'P{i}: [{p.text}] align={p.alignment}')
    if p.text.strip():
        for ri, r in enumerate(p.runs):
            flags = []
            if r.bold: flags.append('BOLD')
            if r.underline: flags.append('UNDERLINE')
            flag_str = f' [{", ".join(flags)}]' if flags else ''
            print(f'  R{ri}: [{r.text}]{flag_str}')
