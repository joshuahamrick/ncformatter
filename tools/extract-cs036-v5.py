import docx

doc = docx.Document(r'c:\Users\jhamrick\Downloads\- Addendum B - CS036 - Verification of Mortgage - Frost - V5.0.docx')

# Get all tables
for ti, table in enumerate(doc.tables):
    print(f'=== Table {ti} ===')
    for ri, row in enumerate(table.rows):
        cells = []
        for ci, cell in enumerate(row.cells):
            text = cell.text.strip()
            cells.append(f'C{ci}:[{text}]')
        print(f'  R{ri}: {" | ".join(cells)}')

# Get paragraphs
print()
print('=== All paragraphs ===')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    bold_runs = [r.text for r in p.runs if r.bold]
    uline_runs = [r.text for r in p.runs if r.underline]
    flags = []
    if bold_runs:
        flags.append('BOLD')
    if uline_runs:
        flags.append('UNDERLINE')
    if p.alignment:
        flags.append(f'ALIGN={p.alignment}')
    flag_str = f' [{", ".join(flags)}]' if flags else ''
    if t:
        print(f'  P{i}: {t}{flag_str}')
