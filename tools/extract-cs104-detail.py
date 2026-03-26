import docx

doc = docx.Document(r'c:\Users\jhamrick\Downloads\CS104 - SCRA Solicitation - Flat Branch - v1.0 (1).docx')

print('=== ALL PARAGRAPHS WITH FULL RUN DETAIL ===')
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if t.strip() or True:  # Show all including empty
        flags = []
        if p.alignment: flags.append(f'ALIGN={p.alignment}')
        flag_str = f' [{", ".join(flags)}]' if flags else ''
        print(f'\nP{i}: [{t}]{flag_str}')
        if t.strip():
            for ri, r in enumerate(p.runs):
                rflags = []
                if r.bold: rflags.append('BOLD')
                if r.underline: rflags.append('UNDERLINE')
                if r.italic: rflags.append('ITALIC')
                if r.font.size: rflags.append(f'SIZE={r.font.size}')
                if r.font.name: rflags.append(f'FONT={r.font.name}')
                if r.font.color and r.font.color.rgb: rflags.append(f'COLOR={r.font.color.rgb}')
                rflag_str = f' [{", ".join(rflags)}]' if rflags else ''
                print(f'  R{ri}: [{r.text}]{rflag_str}')

# Check for hyperlinks
print('\n\n=== HYPERLINKS ===')
for i, p in enumerate(doc.paragraphs):
    for rel in p.part.rels.values():
        if 'hyperlink' in rel.reltype:
            pass
    # Check XML for hyperlinks
    xml = p._element.xml
    if 'hyperlink' in xml:
        print(f'P{i}: Has hyperlink in XML')
        # Extract href
        import re
        links = re.findall(r'r:id="(rId\d+)"', xml)
        for link_id in links:
            if link_id in p.part.rels:
                rel = p.part.rels[link_id]
                print(f'  Link: {rel.target_ref}')

# Check for tables
print('\n\n=== TABLES ===')
for ti, table in enumerate(doc.tables):
    print(f'Table {ti}:')
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if cell.text.strip():
                print(f'  R{ri}C{ci}: [{cell.text.strip()}]')
