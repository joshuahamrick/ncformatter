#!/usr/bin/env python3
"""Generate SI002 template using API modules directly"""

import sys
import os
import json
from pathlib import Path

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document

# Import the IR building function
import importlib.util
spec = importlib.util.spec_from_file_location("process_doc", Path(__file__).parent.parent / "api" / "process-doc.py")
process_doc_module = importlib.util.module_from_spec(spec)
spec.loader.exec_module(process_doc_module)

# Import generate template functions
spec2 = importlib.util.spec_from_file_location("generate_template", Path(__file__).parent.parent / "api" / "generate-template.py")
generate_template_module = importlib.util.module_from_spec(spec2)
spec2.loader.exec_module(generate_template_module)

def main():
    doc_path = Path(__file__).parent.parent / "SI002 - SII Document Request - Triad - V1.0.docx"
    output_path = Path(__file__).parent.parent / "formatter examples" / "SI002" / "SI002-formatted.html"
    
    if not doc_path.exists():
        print(f"ERROR: Document not found: {doc_path}")
        sys.exit(1)
    
    # Load document
    print(f"Loading document: {doc_path}")
    doc = Document(str(doc_path))
    print(f"Document loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    
    # Build IR
    print("Building IR...")
    ir = process_doc_module._build_ir_document(doc)
    print(f"IR built: {len(ir['blocks'])} blocks")
    
    # Check for OpenAI API key
    api_key = os.environ.get('OPENAI_API_KEY')
    if not api_key:
        print("ERROR: OPENAI_API_KEY environment variable not set")
        print("Please set it before running this script")
        sys.exit(1)
    
    # Initialize OpenAI client
    import openai
    client = openai.OpenAI(api_key=api_key)
    
    # Load few-shot examples
    print("Loading few-shot examples...")
    few_shot_examples = generate_template_module.load_few_shot_examples()
    print(f"Loaded {len(few_shot_examples)} examples")
    
    # Build prompt
    print("Building prompt...")
    system_prompt, user_message, few_shot_text = generate_template_module.build_prompt(ir, few_shot_examples)
    full_system_prompt = system_prompt + "\n\n" + few_shot_text
    
    # Call OpenAI
    print("Calling OpenAI API...")
    print(f"System prompt length: {len(full_system_prompt)}")
    print(f"User message length: {len(user_message)}")
    
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": full_system_prompt},
            {"role": "user", "content": user_message}
        ],
        temperature=0,
        max_tokens=8000
    )
    
    html = response.choices[0].message.content.strip()
    print(f"Generated HTML length: {len(html)}")
    
    # Remove markdown code blocks if present
    if html.startswith('```html'):
        html = html.replace('```html', '').replace('```', '').strip()
    elif html.startswith('```'):
        html = html.replace('```', '').strip()
    
    # Normalize HTML
    html = generate_template_module.normalize_html(html)
    
    # Ensure output directory exists
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Save output
    print(f"Saving to: {output_path}")
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)
    
    print("Done!")

if __name__ == '__main__':
    main()

